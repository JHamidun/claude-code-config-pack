#!/usr/bin/env python3
"""Упаковка SEO-материалов в Word (.docx) для маркетолога/стейкхолдера.

Берёт markdown-файлы (бриф, кластер, статья-черновик) и собирает один аккуратный
.docx: обложка + нативные стили Word (заголовки, таблицы, списки), мета статьи из
YAML-frontmatter. Универсально: подаёшь любые md по фазам машины.

CLI:
    python build_report_docx.py --title "Бренд — SEO" --out report.docx \
        --brief research/brief-x.md --cluster research/cluster-x.md --draft drafts/x.md
Любой из --brief/--cluster/--draft необязателен. Порядок в документе: brief → cluster → draft.

Требует: pip install python-docx
"""
import argparse
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("pip install python-docx")

# Нейтральный акцентный цвет; поменяйте под свой бренд.
BRAND = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x66, 0x66, 0x66)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def _inline(p, text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)        # ссылка -> текст
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part:
            p.add_run(part)


def _table(doc, headers, rows, fill="4A4A4A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(htxt); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(10)
        _shade(c, fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _inline(cells[i].paragraphs[0], str(val))
            for rr in cells[i].paragraphs[0].runs:
                rr.font.size = Pt(10)


def render_markdown(doc, body):
    """markdown -> docx (заголовки, таблицы, списки, абзацы, **bold**, ссылки->текст)."""
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i]
        if re.match(r"^#\s+", ln):
            h = doc.add_heading(level=1); r = h.add_run(re.sub(r"^#\s+", "", ln)); r.font.color.rgb = DARK
        elif re.match(r"^##\s+", ln):
            h = doc.add_heading(level=2); r = h.add_run(re.sub(r"^##\s+", "", ln)); r.font.color.rgb = BRAND
        elif re.match(r"^###\s+", ln):
            h = doc.add_heading(level=3); r = h.add_run(re.sub(r"^###\s+", "", ln)); r.font.color.rgb = DARK
        elif ln.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = []
            for rrow in block:
                cells = [c.strip() for c in rrow.strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):
                    rows.append(cells)
            if rows:
                _table(doc, rows[0], rows[1:])
            continue
        elif re.match(r"^\s*[-*]\s+", ln):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                _inline(p, re.sub(r"^\s*[-*]\s+", "", lines[i])); i += 1
            continue
        elif re.match(r"^\s*\d+\.\s+", ln):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                p = doc.add_paragraph(style="List Number")
                _inline(p, re.sub(r"^\s*\d+\.\s+", "", lines[i])); i += 1
            continue
        elif ln.strip().startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(re.sub(r"^>\s*", "", ln).strip()); r.italic = True; r.font.color.rgb = GREY
        elif ln.strip():
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
            _inline(p, ln.strip())
        i += 1


def add_md_file(doc, path, meta_block=False):
    with open(path, encoding="utf-8") as f:
        raw = f.read()
    m = re.match(r"^---\n(.*?)\n---\n", raw, re.S)
    body = raw
    if m:
        fm = {}
        for line in m.group(1).splitlines():
            if ":" in line:
                k, v = line.split(":", 1); fm[k.strip()] = v.strip().strip('"')
        body = raw[m.end():]
        if meta_block:
            for key in ("title", "description"):
                if fm.get(key):
                    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
                    p.add_run(key.capitalize() + ": ").bold = True
                    p.add_run(fm[key]).font.size = Pt(10)
    render_markdown(doc, body)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--title", required=True)
    ap.add_argument("--subtitle", default="Семантическое ядро + SEO/AEO-материал")
    ap.add_argument("--out", required=True)
    ap.add_argument("--brief"); ap.add_argument("--cluster"); ap.add_argument("--draft")
    args = ap.parse_args()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(args.title); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BRAND
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(args.subtitle); r.font.size = Pt(15); r.font.color.rgb = DARK

    if args.brief:
        doc.add_page_break(); add_md_file(doc, args.brief)
    if args.cluster:
        doc.add_page_break(); add_md_file(doc, args.cluster)
    if args.draft:
        doc.add_page_break(); add_md_file(doc, args.draft, meta_block=True)

    doc.save(args.out)
    print("OK:", args.out)


if __name__ == "__main__":
    main()
