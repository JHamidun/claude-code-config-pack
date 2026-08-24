#!/usr/bin/env python3
"""Reusable .docx builder for habr-post pipeline.

Reads FINAL.md + cover.jpg from a working directory, produces <slug>-FINAL.docx
with the standard formatting used across the article series:
- Calibri 11pt body, NAVY (#1A3B6B) headings, 2.5cm margins
- Cover image centered at top, 16cm wide
- Markdown tables → Word tables with header shading
- Code blocks → grey-shaded single-cell tables with Consolas 10pt
- Blockquotes → italic with left border
- Hyperlinks preserved

Usage:
    python build_docx.py --workdir ./article-myslug \\
        --md FINAL.md --cover cover.jpg --slug myslug-FINAL
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

COLOR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_NAVY = RGBColor(0x1A, 0x3B, 0x6B)
COLOR_GREY = RGBColor(0x7A, 0x7A, 0x7A)
COLOR_LINK = RGBColor(0x2A, 0x6F, 0xDB)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2A6FDB")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|\[[^\]]+\]\([^)]+\))"
)


def add_inline_runs(paragraph, text, base_size=Pt(11)):
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            r.font.size = base_size
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            r.font.size = base_size
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            r.font.size = base_size
        elif token.startswith("["):
            mlnk = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mlnk:
                add_hyperlink(paragraph, mlnk.group(2), mlnk.group(1))
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = base_size


def style_heading(para, level):
    para.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.keep_with_next = True
    for run in para.runs:
        run.bold = True
        run.font.color.rgb = COLOR_NAVY
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.allow_autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, "F2F2F2")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    for i, line in enumerate(code_text.splitlines() or [""]):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_BLACK
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_blockquote(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "B0B0B0")
        pBdr.append(left)
        pPr.append(pBdr)
        if line.strip():
            add_inline_runs(p, line)
            for r in p.runs:
                r.italic = True
                if not r.font.color.rgb:
                    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)


def parse_table(table_lines):
    rows = []
    for ln in table_lines:
        ln = ln.strip()
        if not ln.startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            content = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, content, base_size=Pt(10))
            if ri == 0:
                set_cell_shading(cell, "E8E8E8")
                for r in p.runs:
                    r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_paragraph_normal(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    add_inline_runs(p, text)


def add_horizontal_rule(doc, color="C0C0C0"):
    hr_p = doc.add_paragraph()
    pPr = hr_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# Fallback byline when FINAL.md has no "**Автор:**" line. Deliberately a placeholder:
# a real name here would silently sign someone else's article.
DEFAULT_AUTHOR = "<Имя Фамилия>, <должность>"


def build(workdir: Path, md_filename: str, cover_filename: str, slug: str,
         default_author: str = DEFAULT_AUTHOR) -> Path:
    workdir = workdir.resolve()
    md_path = workdir / md_filename
    cover_path = workdir / cover_filename
    out_path = workdir / f"{slug}.docx"

    if not md_path.exists():
        raise FileNotFoundError(f"FINAL md not found: {md_path}")

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    title = None
    meta = {}
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("# ") and title is None:
            title = ln[2:].strip()
            i += 1
            break
        i += 1

    while i < len(lines):
        ln = lines[i].strip()
        if ln == "---":
            i += 1
            break
        if ln.startswith("**") and "**" in ln[2:]:
            m = re.match(r"\*\*([^*]+?):\*\*\s*(.*)", ln)
            if m:
                meta[m.group(1).strip()] = m.group(2).strip()
        i += 1

    doc = Document()
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if cover_path.exists():
        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_para.add_run()
        cover_run.add_picture(str(cover_path), width=Cm(16))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    tr = title_p.add_run(title or "")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = COLOR_NAVY

    author = meta.get("Автор", default_author)
    byline_p = doc.add_paragraph()
    byline_p.paragraph_format.space_after = Pt(2)
    br = byline_p.add_run(author)
    br.italic = True
    br.font.size = Pt(11)
    br.font.color.rgb = COLOR_GREY

    meta_bits = []
    for key in ("Площадка", "Формат", "Статус", "Дата"):
        if meta.get(key):
            meta_bits.append(f"{key}: {meta[key]}")
    if meta_bits:
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(8)
        mr = meta_p.add_run(" · ".join(meta_bits))
        mr.font.size = Pt(9)
        mr.font.color.rgb = COLOR_GREY

    add_horizontal_rule(doc, color="B0B0B0")

    body_lines = lines[i:]
    ix = 0
    while ix < len(body_lines):
        ln = body_lines[ix]
        stripped = ln.strip()

        if stripped == "---":
            add_horizontal_rule(doc)
            ix += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(stripped[4:].strip())
            style_heading(p, 3)
            ix += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.add_run(stripped[3:].strip())
            style_heading(p, 2)
            ix += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.add_run(stripped[2:].strip())
            style_heading(p, 1)
            ix += 1
            continue

        if stripped.startswith("```"):
            ix += 1
            code_lines = []
            while ix < len(body_lines) and not body_lines[ix].strip().startswith("```"):
                code_lines.append(body_lines[ix])
                ix += 1
            ix += 1
            add_code_block(doc, "\n".join(code_lines))
            continue

        if stripped.startswith(">"):
            bq_lines = []
            while ix < len(body_lines) and body_lines[ix].strip().startswith(">"):
                content = body_lines[ix].strip().lstrip(">").lstrip()
                bq_lines.append(content)
                ix += 1
            add_blockquote(doc, bq_lines)
            continue

        if (stripped.startswith("|")
                and ix + 1 < len(body_lines)
                and body_lines[ix + 1].strip().startswith("|")
                and "---" in body_lines[ix + 1]):
            tbl_lines = []
            while ix < len(body_lines) and body_lines[ix].strip().startswith("|"):
                tbl_lines.append(body_lines[ix])
                ix += 1
            rows = parse_table(tbl_lines)
            add_table(doc, rows)
            continue

        if (stripped.startswith("- ") or stripped.startswith("* ")) and stripped[1] == " ":
            add_bullet(doc, stripped[2:].strip())
            ix += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m_num:
            add_numbered(doc, m_num.group(2).strip())
            ix += 1
            continue

        if not stripped:
            ix += 1
            continue

        para_lines = [stripped]
        ix += 1
        while ix < len(body_lines):
            nxt = body_lines[ix].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "```", ">", "|", "- ", "* ")) or re.match(r"^\d+\.\s+", nxt):
                break
            if nxt == "---":
                break
            para_lines.append(nxt)
            ix += 1
        add_paragraph_normal(doc, " ".join(para_lines))

    saved_path = out_path
    try:
        doc.save(str(saved_path))
    except PermissionError:
        saved_path = workdir / f"{slug}.v2.docx"
        doc.save(str(saved_path))
        print(f"NOTE: {out_path.name} is locked. Wrote to {saved_path.name} instead.")
    print(f"Saved: {saved_path} ({saved_path.stat().st_size} bytes)")

    test = Document(str(saved_path))
    print(f"Validation: opened OK. Paragraphs={len(test.paragraphs)}, "
          f"Tables={len(test.tables)}, Sections={len(test.sections)}")

    parts = text.split("---", 2)
    body_text = parts[2] if len(parts) >= 3 else text
    words = re.findall(r"\S+", re.sub(r"```[\s\S]*?```", "", body_text))
    print(f"Word count (body, excluding code blocks): {len(words)}")

    return saved_path


def main():
    parser = argparse.ArgumentParser(description="Build .docx from FINAL.md + cover")
    parser.add_argument("--workdir", required=True, type=Path,
                        help="Directory with FINAL.md and cover image")
    parser.add_argument("--md", default="FINAL.md", help="Markdown filename")
    parser.add_argument("--cover", default="cover.jpg", help="Cover image filename")
    parser.add_argument("--slug", required=True, help="Output slug (without .docx)")
    parser.add_argument("--author", default=DEFAULT_AUTHOR,
                        help="Fallback author byline if FINAL.md has no metadata block. "
                             "Normally comes from the article header; take it from "
                             "~/.claude/author-profile.md")
    args = parser.parse_args()

    try:
        build(args.workdir, args.md, args.cover, args.slug, args.author)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
