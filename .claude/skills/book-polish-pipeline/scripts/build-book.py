# -*- coding: utf-8 -*-
"""
build-book.py — собирает EPUB + DOCX из глав книги.

Picks highest-priority version per chapter (v13.2 > v13 > v12 > ... > FINAL > proofread > DRAFT).
Strips frontmatter, source-trace HTML comments, EDIT-NOTES, PROOFREAD-NOTES.

Usage:
    BOOK_ROOT=./book python build-book.py [--version v13.2]

Конфигурация книги (название, автор, издатель, состав глав) НЕ зашита в скрипт —
она лежит в `$BOOK_ROOT/book.json`. Образец с пояснениями каждого поля:
`../templates/book.example.json` — скопируй его в корень книги и заполни своим.

Environment:
    BOOK_ROOT              — путь к корню книги (обязателен; папка с chapters/ и book.json)
    BOOK_OUT               — путь к dist (default: BOOK_ROOT/dist)
    BOOK_OUTPUT_NAME       — имя выходного файла без расширения (default: из book.json)
    BOOK_VERSION_PREFERENCE — переопределить порядок версий (comma-separated)

Зависимости: pip install ebooklib python-docx
"""
import io, os, sys, re, html, json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')


# Работа — в main(), под `if __name__ == "__main__"`. На верхнем уровне модуля
# только определения: импорт этого файла (линтер с исполнением, автодополнение
# в редакторе, `python -c "import ..."`) не должен ничего запускать и писать.
def main():
    if '--help' in sys.argv or '-h' in sys.argv:
        # --help must not touch the filesystem or build anything: print the module
        # docstring and exit before any work starts.
        print(__doc__ or 'no docstring')
        sys.exit(0)

    # CLI parsing — accept --version v13.2 to override output filename
    _cli_version = None
    for i, arg in enumerate(sys.argv[1:]):
        if arg == '--version' and i + 1 < len(sys.argv) - 1:
            _cli_version = sys.argv[i + 2]
            break
        if arg.startswith('--version='):
            _cli_version = arg.split('=', 1)[1]
            break

    _book_root_env = os.environ.get('BOOK_ROOT')
    if not _book_root_env:
        sys.exit(
            'BOOK_ROOT не задан.\n'
            'Это путь к корню книги — папке, внутри которой лежат chapters/ и book.json.\n'
            '  Windows:  set BOOK_ROOT=C:\\path\\to\\book\n'
            '  bash:     export BOOK_ROOT=./book\n'
            'Образец book.json — ../templates/book.example.json'
        )
    BOOK_ROOT = Path(_book_root_env).expanduser()
    if not BOOK_ROOT.is_dir():
        sys.exit('BOOK_ROOT указывает на несуществующую папку: %s' % BOOK_ROOT)

    # === Конфигурация книги — из $BOOK_ROOT/book.json, а не из кода ===
    # Скрипт общий для любой книги; всё, что относится к конкретной книге
    # (название, автор, издатель, порядок глав), живёт рядом с самой книгой.
    CONFIG_PATH = BOOK_ROOT / 'book.json'
    if not CONFIG_PATH.exists():
        sys.exit(
            'Нет файла %s.\n'
            'Скопируй образец и заполни своим:\n'
            '  cp ../templates/book.example.json "%s"' % (CONFIG_PATH, CONFIG_PATH)
        )
    with io.open(CONFIG_PATH, encoding='utf-8') as _f:
        CFG = json.load(_f)

    _missing = [k for k in ('title', 'author', 'chapters') if not CFG.get(k)]
    if _missing:
        sys.exit('В %s не заполнены обязательные поля: %s' % (CONFIG_PATH, ', '.join(_missing)))

    # === Проверка зависимостей ДО сборки ===
    # `from ebooklib import epub` и `import markdown` живут внутри build_epub(), а тот
    # вызывается в самом конце — после того как все главы прочитаны и обработаны.
    # На машине, где зависимости ставились строго по requirements.txt (ни ebooklib, ни
    # markdown там раньше не значились), человек получал ModuleNotFoundError на
    # последней строке многоминутной сборки. Проверяем здесь: до первой главы.
    _deps = [
        ('ebooklib', 'ebooklib', 'EPUB'),
        ('markdown', 'markdown', 'EPUB'),
        ('docx', 'python-docx', 'DOCX'),
    ]
    _absent = []
    for _mod, _pip, _fmt in _deps:
        try:
            __import__(_mod)
        except ImportError:
            _absent.append((_mod, _pip, _fmt))
    if _absent:
        sys.exit(
            'Не хватает зависимостей для сборки книги:\n'
            + '\n'.join('  • %s — нужен для %s (pip install %s)' % (m, f, p)
                        for m, p, f in _absent)
            + '\n  Одной строкой: pip install '
            + ' '.join(p for _, p, _ in _absent)
            + '\n  Все три перечислены в requirements.txt пака.'
        )

    OUT_DIR = Path(os.environ.get('BOOK_OUT', str(BOOK_ROOT / 'dist')))
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_NAME = os.environ.get('BOOK_OUTPUT_NAME') or CFG.get('output_name') or 'book'
    VERSION_TAG = _cli_version or os.environ.get('BOOK_VERSION_TAG', 'v1')

    # === Chapter order with display titles ===
    # (slug, part_title, chapter_title, illustration_filename_or_None)
    # В book.json это массив массивов из 4 элементов; null там, где элемента нет.
    CHAPTERS = [tuple((list(c) + [None, None, None, None])[:4]) for c in CFG['chapters']]

    ILLUSTRATIONS_DIR = OUT_DIR / 'illustrations'
    COVER_PATH = OUT_DIR / (CFG.get('cover') or 'cover/cover.png')

    BOOK_TITLE = CFG['title']
    BOOK_SUBTITLE = CFG.get('subtitle', '')
    BOOK_COVER_QUOTE = CFG.get('cover_quote', '')
    BOOK_AUTHOR = CFG['author']
    BOOK_LANG = CFG.get('lang', 'ru')
    BOOK_PUBLISHER = CFG.get('publisher', '')
    BOOK_YEAR = str(CFG.get('year', ''))
    BOOK_DESCRIPTION = CFG.get('description', '')

    # ============================================================
    # CLEANUP — strip technical content from markdown
    # ============================================================
    def clean_chapter_md(text):
        """Remove frontmatter, source-trace, EDIT-NOTES, PROOFREAD-NOTES blocks."""
        # 1. Strip YAML frontmatter at start
        text = re.sub(r'^---\n.*?\n---\n+', '', text, count=1, flags=re.DOTALL)

        # 2. Strip <!-- source-trace: ... --> blocks
        text = re.sub(r'<!--\s*source-trace:.*?-->\s*', '', text, flags=re.DOTALL)

        # 3. Strip "EDIT-NOTES (не для печати):" through end-of-fence
        text = re.sub(r'\n+EDIT-NOTES.*?(?=\n```|\Z)', '', text, flags=re.DOTALL)

        # 4. Strip ```PROOFREAD-NOTES:...``` code fences (greedy, multiple variants)
        text = re.sub(r'```\s*\n?PROOFREAD-NOTES:.*?\n```', '', text, flags=re.DOTALL)
        text = re.sub(r'```[^\n]*\n?PROOFREAD-NOTES:.*?```', '', text, flags=re.DOTALL)
        text = re.sub(r'PROOFREAD-NOTES:[\s\S]*?(?=\n```|\n##\s|\Z)', '', text)
        text = re.sub(r'PROOFREAD NOTES:[\s\S]*?(?=\n```|\n##\s|\Z)', '', text)
        # Also strip the editor metadata blocks at end of chapter
        text = re.sub(r'\n+EDIT[\-\s]?NOTES?:[\s\S]*?(?=\n```|\Z)', '', text)
        # And bare lines like "- 0 исправлений орфографии" that may leak
        text = re.sub(r'^\s*-\s*\d+\s+исправлен[а-я]+.*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*•.*типографик.*$', '', text, flags=re.MULTILINE)

        # 5. Strip leftover empty code fences
        text = re.sub(r'```\s*\n```', '', text)
        text = re.sub(r'```\s*\Z', '', text)
        text = re.sub(r'\n```\s*$', '', text, flags=re.MULTILINE)

        # 6. Strip first H1 (we'll set our own from CHAPTERS table)
        text = re.sub(r'^#\s+[^\n]+\n+', '', text.lstrip(), count=1)

        # 7. Collapse 3+ blank lines to 2
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()


    # Version preference — from newest to oldest. Picks first existing file per chapter.
    # Override via env var: BOOK_VERSION_PREFERENCE="v13.2,v13.1,v13,v12,FINAL,proofread,voice-pass,DRAFT"
    DEFAULT_VERSION_PREFERENCE = [
        'v24', 'v23', 'v22',          # newest versions (added for v39 build)
        'v21', 'v20', 'v19', 'v18', 'v17', 'v16', 'v15',
        'v14',
        'v13.2', 'v13.1', 'v13',
        'v12', 'v11', 'v10', 'v9', 'v8', 'v1',
        'FINAL',        # post-book-post FINAL.md
        'proofread',    # DRAFT.proofread.md
        'voice-pass',   # DRAFT.voice-pass.md
        'DRAFT',        # initial DRAFT.md
    ]

    VERSION_PREFERENCE = os.environ.get('BOOK_VERSION_PREFERENCE', '').split(',') if os.environ.get('BOOK_VERSION_PREFERENCE') else DEFAULT_VERSION_PREFERENCE
    VERSION_PREFERENCE = [v.strip() for v in VERSION_PREFERENCE if v.strip()]


    def load_chapter(slug):
        """Find the highest-priority version of the chapter and load+clean it."""
        ch_dir = BOOK_ROOT / 'chapters' / slug
        for ver in VERSION_PREFERENCE:
            # FINAL.md is special — no DRAFT prefix
            if ver == 'FINAL':
                p = ch_dir / 'FINAL.md'
            else:
                p = ch_dir / f'DRAFT.{ver}.md'
            if p.exists():
                return clean_chapter_md(p.read_text(encoding='utf-8'))
        return None


    # ============================================================
    # EPUB BUILDER
    # ============================================================
    def build_epub():
        from ebooklib import epub
        import markdown as md_lib

        book = epub.EpubBook()
        book.set_identifier(OUTPUT_NAME)
        book.set_title(BOOK_TITLE)
        book.set_language(BOOK_LANG)
        book.add_author(BOOK_AUTHOR)
        book.add_metadata('DC', 'publisher', BOOK_PUBLISHER)
        if BOOK_DESCRIPTION:
            book.add_metadata('DC', 'description', BOOK_DESCRIPTION)

        # CSS — book-style typography
        css = '''
    @namespace epub "http://www.idpf.org/2007/ops";
    body { font-family: Georgia, "Times New Roman", serif; line-height: 1.6; margin: 0 1.2em; color: #1a1a1a; }
    h1 { font-size: 1.6em; line-height: 1.2; margin: 1.5em 0 0.8em 0; page-break-before: always; }
    h2 { font-size: 1.2em; margin: 1.6em 0 0.6em 0; line-height: 1.3; }
    h3 { font-size: 1.05em; margin: 1.2em 0 0.5em 0; }
    p  { margin: 0 0 0.7em 0; text-align: justify; text-indent: 1.2em; }
    p:first-of-type, h1+p, h2+p, h3+p, blockquote+p { text-indent: 0; }
    blockquote { border-left: 3px solid #888; margin: 1em 0.5em; padding: 0 0 0 0.8em; font-style: italic; color: #444; }
    ul, ol { margin: 0.5em 0 1em 1.5em; }
    li { margin-bottom: 0.3em; }
    strong { font-weight: 700; }
    em { font-style: italic; }
    hr { border: none; border-top: 1px solid #aaa; margin: 2em 25%; }
    table { border-collapse: collapse; margin: 1em 0; width: 100%; }
    th, td { border: 1px solid #888; padding: 0.4em 0.6em; text-align: left; }
    th { background: #f0f0f0; font-weight: 700; }
    code { font-family: Consolas, monospace; font-size: 0.9em; background: #f4f4f4; padding: 0 0.2em; }
    .part-title { text-align: center; page-break-before: always; padding-top: 30%; font-size: 2em; font-weight: 300; letter-spacing: 0.05em; }
    .chapter-num { color: #888; font-size: 0.7em; font-weight: 400; display: block; margin-bottom: 0.5em; letter-spacing: 0.1em; text-transform: uppercase; }
    .title-page { text-align: center; padding-top: 12%; }
    .title-page .cover-quote { font-size: 1.05em; font-style: italic; color: #444; max-width: 80%; margin: 0 auto 2.5em auto; line-height: 1.45; }
    .title-page h1 { font-size: 2.4em; margin: 0.3em 0; page-break-before: auto; letter-spacing: 0.02em; text-transform: uppercase; font-weight: 700; }
    .title-page .cover-subtitle { font-size: 1.1em; color: #555; margin-top: 1.5em; max-width: 75%; margin-left: auto; margin-right: auto; line-height: 1.4; }
    .title-page .author { font-size: 1.3em; margin-top: 3em; color: #444; }
    .title-page .publisher { font-size: 1em; margin-top: 4em; color: #888; }
    '''
        css_item = epub.EpubItem(uid='style_book', file_name='style/book.css', media_type='text/css', content=css)
        book.add_item(css_item)

        md_converter = md_lib.Markdown(extensions=['extra', 'sane_lists'])

        spine = ['nav']
        toc = []

        # Title page (cover)
        title_html = f'''<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
    <head><title>{BOOK_TITLE}</title><link rel="stylesheet" href="style/book.css" type="text/css"/></head>
    <body><div class="title-page">
    <p class="cover-quote">{BOOK_COVER_QUOTE}</p>
    <h1>{BOOK_TITLE}</h1>
    <p class="cover-subtitle">{BOOK_SUBTITLE}</p>
    <p class="author">{BOOK_AUTHOR}</p>
    <p class="publisher">{' · '.join(x for x in (BOOK_PUBLISHER, BOOK_YEAR) if x)}</p>
    </div></body></html>'''
        title_item = epub.EpubHtml(title='Титул', file_name='000-title.xhtml', lang=BOOK_LANG)
        title_item.content = title_html
        title_item.add_item(css_item)
        book.add_item(title_item)
        spine.append(title_item)

        # Cover image
        if COVER_PATH.exists():
            with open(COVER_PATH, 'rb') as f:
                cover_data = f.read()
            book.set_cover('cover.png', cover_data)
            print(f'[cover] embedded: {COVER_PATH.name} ({len(cover_data)//1024} KB)')

        # Chapters
        for idx, item in enumerate(CHAPTERS, 1):
            slug, part_title, ch_title = item[0], item[1], item[2]
            illust_filename = item[3] if len(item) > 3 else None
            body_md = load_chapter(slug)
            if body_md is None:
                print(f'[skip] {slug}: no DRAFT.proofread.md', file=sys.stderr)
                continue
            body_html = md_converter.convert(body_md)
            md_converter.reset()

            # Insert illustration <img> after H1 if present
            illust_html = ''
            if illust_filename:
                ill_path = ILLUSTRATIONS_DIR / illust_filename
                if ill_path.exists():
                    with open(ill_path, 'rb') as f:
                        ill_bytes = f.read()
                    ill_item = epub.EpubItem(
                        uid=f'img_{idx}',
                        file_name=f'images/{illust_filename}',
                        media_type='image/png',
                        content=ill_bytes,
                    )
                    book.add_item(ill_item)
                    illust_html = f'<div style="text-align:center; margin: 1em 0 1.5em 0;"><img src="images/{illust_filename}" style="max-width:85%; height:auto;" alt=""/></div>'

            sections_html = ''

            if part_title:
                # Standalone part-title page
                part_safe = html.escape(part_title)
                part_html = f'''<html xmlns="http://www.w3.org/1999/xhtml"><head><title>{part_safe}</title><link rel="stylesheet" href="style/book.css" type="text/css"/></head><body><div class="part-title">{part_safe}</div></body></html>'''
                part_item = epub.EpubHtml(title=part_title, file_name=f'{idx:03d}a-part.xhtml', lang=BOOK_LANG)
                part_item.content = part_html
                part_item.add_item(css_item)
                book.add_item(part_item)
                spine.append(part_item)
                toc.append(epub.Link(part_item.file_name, part_title, f'part-{idx}'))

            ch_safe = html.escape(ch_title)
            # Split chapter title into "Глава N." chapter-num + main title
            m = re.match(r'^(Глава\s+\d+\.)\s+(.+)$', ch_title)
            if m:
                ch_num, ch_main = html.escape(m.group(1)), html.escape(m.group(2))
                h1 = f'<h1><span class="chapter-num">{ch_num}</span>{ch_main}</h1>'
            else:
                h1 = f'<h1>{ch_safe}</h1>'

            chapter_html = f'''<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
    <head><title>{ch_safe}</title><link rel="stylesheet" href="style/book.css" type="text/css"/></head>
    <body>
    {h1}
    {illust_html}
    {body_html}
    </body></html>'''
            ch_item = epub.EpubHtml(title=ch_title, file_name=f'{idx:03d}-{slug}.xhtml', lang=BOOK_LANG)
            ch_item.content = chapter_html
            ch_item.add_item(css_item)
            book.add_item(ch_item)
            spine.append(ch_item)
            toc.append(epub.Link(ch_item.file_name, ch_title, f'ch-{slug}'))
            print(f'[ok] ch{idx:02d}: {ch_title} ({len(body_md)} chars)')

        book.toc = tuple(toc)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = spine

        out_path = OUT_DIR / f'{OUTPUT_NAME}-{VERSION_TAG}.epub'
        epub.write_epub(str(out_path), book, {})
        print(f'\n[EPUB] {out_path} ({out_path.stat().st_size // 1024} KB)')
        return out_path


    # ============================================================
    # DOCX BUILDER
    # ============================================================
    def build_docx():
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Page setup A5-ish for book feel
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Customize Normal style
        normal = doc.styles['Normal']
        normal.font.name = 'Georgia'
        normal.font.size = Pt(11.5)
        normal.paragraph_format.line_spacing = 1.4
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.first_line_indent = Cm(0.7)

        # Heading 1 - chapter
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Georgia'
        h1.font.size = Pt(20)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h1.paragraph_format.space_before = Pt(36)
        h1.paragraph_format.space_after = Pt(18)
        h1.paragraph_format.page_break_before = True

        # Heading 2 - section
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Georgia'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)

        # Heading 3
        h3 = doc.styles['Heading 3']
        h3.font.name = 'Georgia'
        h3.font.size = Pt(12)
        h3.font.bold = True

        # Cover image as the very first page
        if COVER_PATH.exists():
            from docx.shared import Inches
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_para.add_run().add_picture(str(COVER_PATH), width=Inches(6.0))
            doc.add_page_break()

        # Title page
        # Cover quote
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(140)
        r = p.add_run(BOOK_COVER_QUOTE)
        r.font.size = Pt(13)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        r.font.name = 'Georgia'

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)
        r = p.add_run(BOOK_TITLE.upper())
        r.font.size = Pt(30)
        r.font.bold = True
        r.font.name = 'Georgia'

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run(BOOK_SUBTITLE)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r.font.name = 'Georgia'

        # Author
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        r = p.add_run(BOOK_AUTHOR)
        r.font.size = Pt(16)
        r.font.name = 'Georgia'

        # Publisher
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(80)
        r = p.add_run(' · '.join(x for x in (BOOK_PUBLISHER, BOOK_YEAR) if x))
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        r.font.name = 'Georgia'

        # Process chapters
        from docx.shared import Inches
        for item in CHAPTERS:
            slug, part_title, ch_title = item[0], item[1], item[2]
            illust_filename = item[3] if len(item) > 3 else None
            body_md = load_chapter(slug)
            if body_md is None:
                continue

            if part_title:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(220)
                r = p.add_run(part_title)
                r.font.size = Pt(26)
                r.font.bold = False
                r.font.name = 'Georgia'

            # Chapter heading
            doc.add_heading(ch_title, level=1)

            # Insert illustration centered if present
            if illust_filename:
                ill_path = ILLUSTRATIONS_DIR / illust_filename
                if ill_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(ill_path), width=Inches(4.5))

            # Convert markdown body line by line
            render_md_to_docx(doc, body_md)

        out_path = OUT_DIR / f'{OUTPUT_NAME}-{VERSION_TAG}.docx'
        doc.save(str(out_path))
        print(f'[DOCX] {out_path} ({out_path.stat().st_size // 1024} KB)')
        return out_path


    def render_md_to_docx(doc, md_text):
        """Simple markdown → docx rendering: paragraphs, headings, blockquotes, lists, bold/italic."""
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Split by blocks (blank-line separated)
        blocks = re.split(r'\n\n+', md_text)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            # Heading
            m = re.match(r'^(#{2,6})\s+(.+)$', block)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                doc.add_heading(_strip_md(text), level=min(level, 4))
                continue

            # Horizontal rule
            if re.match(r'^\s*-{3,}\s*$', block) or re.match(r'^\s*\*{3,}\s*$', block):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                r = p.add_run('* * *')
                r.font.size = Pt(11)
                continue

            # Blockquote
            if block.startswith('>'):
                quote_lines = [re.sub(r'^>\s?', '', ln) for ln in block.split('\n')]
                quote_text = ' '.join(quote_lines).strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.right_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                _add_inline_runs(p, quote_text, italic_default=True)
                continue

            # Code block (skip rendering as monospace block, render as plain paragraph)
            if block.startswith('```'):
                inner = re.sub(r'^```[^\n]*\n?', '', block)
                inner = re.sub(r'\n?```\s*$', '', inner)
                for ln in inner.split('\n'):
                    if ln.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(0)
                        r = p.add_run(ln)
                        r.font.name = 'Consolas'
                        r.font.size = Pt(10)
                continue

            # List
            if re.match(r'^\s*[-*+]\s+', block) or re.match(r'^\s*\d+\.\s+', block):
                for ln in block.split('\n'):
                    ln = ln.strip()
                    m = re.match(r'^[-*+]\s+(.+)$', ln) or re.match(r'^\d+\.\s+(.+)$', ln)
                    if m:
                        p = doc.add_paragraph(style='List Bullet' if not re.match(r'^\d', ln) else 'List Number')
                        p.paragraph_format.first_line_indent = Cm(0)
                        _add_inline_runs(p, m.group(1))
                continue

            # Table — rough: detect "| header | header |" then "| --- | --- |"
            if '\n' in block and block.startswith('|'):
                lines = [ln for ln in block.split('\n') if ln.strip().startswith('|')]
                if len(lines) >= 2 and re.match(r'^\|[\s\-:|]+\|$', lines[1].strip()):
                    rows = []
                    for ln in lines:
                        if re.match(r'^\|[\s\-:|]+\|$', ln.strip()):
                            continue
                        cells = [c.strip() for c in ln.strip('|').split('|')]
                        rows.append(cells)
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1' if 'Light Grid Accent 1' in [s.name for s in doc.styles] else 'Table Grid'
                        for i, row_cells in enumerate(rows):
                            for j, cell_text in enumerate(row_cells):
                                if j < len(table.rows[i].cells):
                                    table.rows[i].cells[j].text = _strip_md(cell_text)
                        continue

            # Plain paragraph
            text = block.replace('\n', ' ')
            p = doc.add_paragraph()
            _add_inline_runs(p, text)


    def _add_inline_runs(p, text, italic_default=False):
        """Parse **bold** and *italic* and add as runs."""
        # Pattern catches **bold**, __bold__, *italic*, _italic_, `code`
        pattern = re.compile(r'(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)')
        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
                r.italic = italic_default
            elif part.startswith('__') and part.endswith('__'):
                r = p.add_run(part[2:-2])
                r.bold = True
                r.italic = italic_default
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                r = p.add_run(part[1:-1])
                r.italic = True
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                r = p.add_run(part[1:-1])
                r.italic = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                from docx.shared import Pt
                r.font.name = 'Consolas'
                r.font.size = Pt(10)
                r.italic = italic_default
            else:
                r = p.add_run(part)
                r.italic = italic_default


    def _strip_md(text):
        """Strip markdown markers for plain rendering (headings, table cells)."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', text)
        text = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return text


    print('=== Building EPUB ===')
    epub_path = build_epub()
    print('\n=== Building DOCX ===')
    docx_path = build_docx()
    print(f'\n✓ Done. Files in: {OUT_DIR}')


if __name__ == "__main__":
    main()
