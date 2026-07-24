"""
Универсальный форматтер отчётов: text / markdown / xlsx / docx / json / csv.

Использование:
    from report_formatter import Report

    r = Report(title='Аналитика Захарова', period='07.04-13.05')
    r.section('Ключевые цифры')
    r.kv('Всего сделок', 261)
    r.kv('WON', 0, comment='0 побед за 36 дней')
    r.table(['Стадия', 'Кол-во'], [['NEW', 15], ['LOSE', 178]])
    r.paragraph('Захаров работает конвейером...')
    print(r.to_text())
    r.to_xlsx('~/Downloads/report.xlsx')
"""
import os
import json
import csv
from datetime import datetime


class Report:
    """Универсальный контейнер для отчёта, экспортируемый в разные форматы."""

    def __init__(self, title: str, period: str = '', source: str = 'Битрикс24 (webhook)', subtitle: str = ''):
        self.title = title
        self.period = period
        self.source = source
        self.subtitle = subtitle
        self.generated = datetime.now().strftime('%Y-%m-%d %H:%M')
        self.blocks = []  # list of (type, data)

    # === API for building the report ===

    def section(self, title: str, level: int = 2):
        """Add section header."""
        self.blocks.append(('section', {'title': title, 'level': level}))
        return self

    def kv(self, key: str, value, comment: str = '', highlight: str = None):
        """Add key-value row. highlight: 'red' | 'green' | 'yellow' | None"""
        self.blocks.append(('kv', {'key': key, 'value': value, 'comment': comment, 'highlight': highlight}))
        return self

    def kv_group(self, rows: list):
        """Add multiple key-value rows at once. rows = [(k, v, comment?), ...]"""
        for r in rows:
            if len(r) == 2:
                self.kv(r[0], r[1])
            elif len(r) >= 3:
                self.kv(r[0], r[1], comment=r[2])
        return self

    def table(self, headers: list, rows: list, name: str = ''):
        """Add table with headers and rows."""
        self.blocks.append(('table', {'headers': headers, 'rows': rows, 'name': name}))
        return self

    def paragraph(self, text: str, style: str = 'normal'):
        """Add paragraph. style: 'normal' | 'note' | 'warning' | 'success'"""
        self.blocks.append(('paragraph', {'text': text, 'style': style}))
        return self

    def bullet(self, items: list, style: str = 'normal'):
        """Add bullet list."""
        self.blocks.append(('bullet', {'items': items, 'style': style}))
        return self

    def separator(self):
        """Add visual separator."""
        self.blocks.append(('separator', {}))
        return self

    def conclusion(self, strong: list = None, weak: list = None, questions: list = None, recs: list = None):
        """Add conclusion block."""
        self.blocks.append(('conclusion', {
            'strong': strong or [],
            'weak': weak or [],
            'questions': questions or [],
            'recs': recs or []
        }))
        return self

    # === Text formatter (моноширинный, для чата и email) ===

    def to_text(self, width: int = 60) -> str:
        """Format as monospaced plain text."""
        out = []
        eq = '═' * width
        dash = '─' * width

        # Header
        out.append(eq)
        out.append(self.title.upper().center(width))
        out.append(eq)
        out.append('')
        if self.period:
            out.append(f'Период: {self.period}')
        if self.subtitle:
            out.append(self.subtitle)
        out.append(f'Источник: {self.source} | Сгенерировано: {self.generated}')
        out.append('')

        for kind, data in self.blocks:
            if kind == 'section':
                out.append(dash)
                out.append(data['title'].upper())
                out.append(dash)
                out.append('')
            elif kind == 'kv':
                v = data['value']
                if isinstance(v, float):
                    v = f'{v:,.0f}' if v >= 100 else f'{v:.2f}'
                elif isinstance(v, int):
                    v = f'{v:,}'.replace(',', ' ')
                line = f'• {data["key"]}: {v}'
                if data.get('comment'):
                    line += f'  — {data["comment"]}'
                out.append(line)
            elif kind == 'table':
                if data.get('name'):
                    out.append('')
                    out.append(data['name'])
                out.append('')
                out.extend(_format_text_table(data['headers'], data['rows']))
                out.append('')
            elif kind == 'paragraph':
                text = data['text']
                marker = {'warning': '⚠  ', 'success': '✓  ', 'note': 'Примечание: ', 'normal': ''}.get(data['style'], '')
                # Wrap text at width
                wrapped = _wrap_text(marker + text, width)
                out.extend(wrapped)
                out.append('')
            elif kind == 'bullet':
                for item in data['items']:
                    out.append(f'• {item}')
                out.append('')
            elif kind == 'separator':
                out.append('')
                out.append(dash)
                out.append('')
            elif kind == 'conclusion':
                if data['strong']:
                    out.append('СИЛЬНЫЕ СТОРОНЫ:')
                    for x in data['strong']:
                        out.append(f'+ {x}')
                    out.append('')
                if data['weak']:
                    out.append('СЛАБЫЕ МЕСТА:')
                    for i, x in enumerate(data['weak'], 1):
                        out.append(f'{i}. {x}')
                    out.append('')
                if data['questions']:
                    out.append('ВОПРОСЫ К РАЗБОРУ:')
                    for i, x in enumerate(data['questions'], 1):
                        out.append(f'{i}. {x}')
                    out.append('')
                if data['recs']:
                    out.append('РЕКОМЕНДАЦИИ:')
                    for i, x in enumerate(data['recs'], 1):
                        out.append(f'{i}. {x}')
                    out.append('')

        out.append(eq)
        return '\n'.join(out)

    # === Markdown formatter ===

    def to_md(self) -> str:
        """Format as markdown."""
        out = []
        out.append(f'# {self.title}')
        if self.period:
            out.append(f'*Период: {self.period}*  ')
        out.append(f'*Источник: {self.source} · {self.generated}*')
        out.append('')

        for kind, data in self.blocks:
            if kind == 'section':
                hashes = '#' * (data.get('level', 2) + 1)
                out.append(f'{hashes} {data["title"]}')
                out.append('')
            elif kind == 'kv':
                v = data['value']
                if isinstance(v, (int, float)):
                    v = f'{v:,}'.replace(',', ' ')
                line = f'- **{data["key"]}**: {v}'
                if data.get('comment'):
                    line += f' — _{data["comment"]}_'
                out.append(line)
            elif kind == 'table':
                if data.get('name'):
                    out.append(f'**{data["name"]}**')
                out.append('')
                out.append('| ' + ' | '.join(str(h) for h in data['headers']) + ' |')
                out.append('|' + '|'.join(['---'] * len(data['headers'])) + '|')
                for row in data['rows']:
                    out.append('| ' + ' | '.join(str(c) for c in row) + ' |')
                out.append('')
            elif kind == 'paragraph':
                emoji = {'warning': '⚠️ ', 'success': '✅ ', 'note': '💡 ', 'normal': ''}.get(data['style'], '')
                out.append(f'{emoji}{data["text"]}')
                out.append('')
            elif kind == 'bullet':
                for item in data['items']:
                    out.append(f'- {item}')
                out.append('')
            elif kind == 'separator':
                out.append('---')
                out.append('')
            elif kind == 'conclusion':
                if data['strong']:
                    out.append('### Сильные стороны')
                    for x in data['strong']:
                        out.append(f'- ✅ {x}')
                    out.append('')
                if data['weak']:
                    out.append('### Слабые места')
                    for i, x in enumerate(data['weak'], 1):
                        out.append(f'{i}. ⚠️ {x}')
                    out.append('')
                if data['questions']:
                    out.append('### Вопросы к разбору')
                    for i, x in enumerate(data['questions'], 1):
                        out.append(f'{i}. {x}')
                    out.append('')
                if data['recs']:
                    out.append('### Рекомендации')
                    for i, x in enumerate(data['recs'], 1):
                        out.append(f'{i}. {x}')
                    out.append('')

        return '\n'.join(out)

    # === JSON export ===

    def to_json(self) -> str:
        """Format as JSON (for programmatic consumption)."""
        return json.dumps({
            'title': self.title,
            'period': self.period,
            'source': self.source,
            'subtitle': self.subtitle,
            'generated': self.generated,
            'blocks': self.blocks,
        }, ensure_ascii=False, indent=2, default=str)

    # === XLSX export ===

    def to_xlsx(self, path: str):
        """Write to Excel file with formatting."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        except ImportError:
            raise ImportError('pip install openpyxl')

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Отчёт'

        # Styles
        title_font = Font(name='Calibri', size=16, bold=True, color='00467F')
        section_font = Font(name='Calibri', size=13, bold=True, color='0070C0')
        header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='00467F', end_color='00467F', fill_type='solid')
        alt_fill = PatternFill(start_color='DCE6F1', end_color='DCE6F1', fill_type='solid')
        red_font = Font(name='Calibri', size=11, color='C62828', bold=True)
        green_font = Font(name='Calibri', size=11, color='2E7D32', bold=True)
        yellow_font = Font(name='Calibri', size=11, color='E65100', bold=True)
        bold_font = Font(name='Calibri', size=11, bold=True)
        thin_border = Border(
            left=Side(style='thin', color='B0B0B0'),
            right=Side(style='thin', color='B0B0B0'),
            top=Side(style='thin', color='B0B0B0'),
            bottom=Side(style='thin', color='B0B0B0'),
        )

        row = 1
        ws.cell(row=row, column=1, value=self.title).font = title_font
        row += 1
        if self.period:
            ws.cell(row=row, column=1, value=f'Период: {self.period}').font = Font(size=10, color='808080')
            row += 1
        ws.cell(row=row, column=1, value=f'{self.source} · {self.generated}').font = Font(size=10, color='808080')
        row += 2

        highlight_map = {'red': red_font, 'green': green_font, 'yellow': yellow_font}

        for kind, data in self.blocks:
            if kind == 'section':
                ws.cell(row=row, column=1, value=data['title']).font = section_font
                row += 2
            elif kind == 'kv':
                ws.cell(row=row, column=1, value=data['key']).font = bold_font
                cv = ws.cell(row=row, column=2, value=data['value'])
                if data.get('highlight') in highlight_map:
                    cv.font = highlight_map[data['highlight']]
                if data.get('comment'):
                    ws.cell(row=row, column=3, value=data['comment']).font = Font(color='808080', italic=True)
                row += 1
            elif kind == 'table':
                if data.get('name'):
                    ws.cell(row=row, column=1, value=data['name']).font = bold_font
                    row += 1
                # Headers
                for i, h in enumerate(data['headers'], 1):
                    c = ws.cell(row=row, column=i, value=h)
                    c.font = header_font
                    c.fill = header_fill
                    c.alignment = Alignment(horizontal='center', vertical='center')
                    c.border = thin_border
                row += 1
                # Rows
                for ri, r_data in enumerate(data['rows']):
                    for i, v in enumerate(r_data, 1):
                        c = ws.cell(row=row, column=i, value=v)
                        c.border = thin_border
                        if ri % 2 == 0:
                            c.fill = alt_fill
                    row += 1
                row += 1
            elif kind == 'paragraph':
                ws.cell(row=row, column=1, value=data['text'])
                row += 1
            elif kind == 'bullet':
                for item in data['items']:
                    ws.cell(row=row, column=1, value=f'• {item}')
                    row += 1
                row += 1
            elif kind == 'separator':
                row += 1
            elif kind == 'conclusion':
                if data['strong']:
                    ws.cell(row=row, column=1, value='СИЛЬНЫЕ СТОРОНЫ').font = section_font
                    row += 1
                    for x in data['strong']:
                        c = ws.cell(row=row, column=1, value=f'+ {x}')
                        c.font = green_font
                        row += 1
                    row += 1
                if data['weak']:
                    ws.cell(row=row, column=1, value='СЛАБЫЕ МЕСТА').font = section_font
                    row += 1
                    for i, x in enumerate(data['weak'], 1):
                        c = ws.cell(row=row, column=1, value=f'{i}. {x}')
                        c.font = red_font
                        row += 1
                    row += 1
                if data['questions']:
                    ws.cell(row=row, column=1, value='ВОПРОСЫ К РАЗБОРУ').font = section_font
                    row += 1
                    for i, x in enumerate(data['questions'], 1):
                        ws.cell(row=row, column=1, value=f'{i}. {x}')
                        row += 1
                    row += 1
                if data['recs']:
                    ws.cell(row=row, column=1, value='РЕКОМЕНДАЦИИ').font = section_font
                    row += 1
                    for i, x in enumerate(data['recs'], 1):
                        ws.cell(row=row, column=1, value=f'{i}. {x}')
                        row += 1
                    row += 1

        # Column widths
        for col, w in [(1, 45), (2, 25), (3, 40), (4, 25), (5, 20), (6, 20), (7, 20), (8, 20)]:
            from openpyxl.utils import get_column_letter
            ws.column_dimensions[get_column_letter(col)].width = w

        path = os.path.expanduser(path)
        os.makedirs(os.path.dirname(path), exist_ok=True) if os.path.dirname(path) else None
        wb.save(path)
        return path

    # === DOCX export ===

    def to_docx(self, path: str):
        """Write to Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError('pip install python-docx')

        doc = Document()

        # Title
        h = doc.add_heading(self.title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{self.period}\n{self.source} · {self.generated}' if self.period else f'{self.source} · {self.generated}')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(128, 128, 128)

        for kind, data in self.blocks:
            if kind == 'section':
                doc.add_heading(data['title'], level=data.get('level', 2))
            elif kind == 'kv':
                p = doc.add_paragraph()
                p.add_run(f'{data["key"]}: ').bold = True
                p.add_run(str(data['value']))
                if data.get('comment'):
                    p.add_run(f'  — {data["comment"]}').italic = True
            elif kind == 'table':
                if data.get('name'):
                    doc.add_paragraph(data['name']).runs[0].bold = True
                t = doc.add_table(rows=1 + len(data['rows']), cols=len(data['headers']))
                t.style = 'Table Grid'
                for i, h in enumerate(data['headers']):
                    c = t.rows[0].cells[i]
                    c.text = str(h)
                    c.paragraphs[0].runs[0].bold = True
                for ri, r_data in enumerate(data['rows'], 1):
                    for i, v in enumerate(r_data):
                        t.rows[ri].cells[i].text = str(v)
                doc.add_paragraph()
            elif kind == 'paragraph':
                doc.add_paragraph(data['text'])
            elif kind == 'bullet':
                for item in data['items']:
                    doc.add_paragraph(item, style='List Bullet')
            elif kind == 'conclusion':
                if data['strong']:
                    doc.add_heading('Сильные стороны', level=2)
                    for x in data['strong']:
                        doc.add_paragraph(x, style='List Bullet')
                if data['weak']:
                    doc.add_heading('Слабые места', level=2)
                    for x in data['weak']:
                        doc.add_paragraph(x, style='List Number')
                if data['questions']:
                    doc.add_heading('Вопросы к разбору', level=2)
                    for x in data['questions']:
                        doc.add_paragraph(x, style='List Number')
                if data['recs']:
                    doc.add_heading('Рекомендации', level=2)
                    for x in data['recs']:
                        doc.add_paragraph(x, style='List Number')

        path = os.path.expanduser(path)
        doc.save(path)
        return path

    # === CSV export (только для таблиц) ===

    def to_csv(self, path: str):
        """Write all tables to CSV (concatenated)."""
        path = os.path.expanduser(path)
        with open(path, 'w', encoding='utf-8-sig', newline='') as f:
            w = csv.writer(f)
            w.writerow([self.title])
            if self.period:
                w.writerow([f'Период: {self.period}'])
            w.writerow([])
            for kind, data in self.blocks:
                if kind == 'section':
                    w.writerow([])
                    w.writerow([data['title']])
                elif kind == 'table':
                    if data.get('name'):
                        w.writerow([data['name']])
                    w.writerow(data['headers'])
                    for row in data['rows']:
                        w.writerow(row)
                    w.writerow([])
                elif kind == 'kv':
                    w.writerow([data['key'], data['value'], data.get('comment', '')])
        return path

    # === Dispatch ===

    def output(self, format: str = 'text', path: str = None) -> str:
        """Output in specified format. Returns text or path."""
        if format == 'text':
            return self.to_text()
        elif format == 'md':
            return self.to_md()
        elif format == 'json':
            return self.to_json()
        elif format == 'xlsx':
            return self.to_xlsx(path or f'~/Downloads/report_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx')
        elif format == 'docx':
            return self.to_docx(path or f'~/Downloads/report_{datetime.now().strftime("%Y%m%d_%H%M")}.docx')
        elif format == 'csv':
            return self.to_csv(path or f'~/Downloads/report_{datetime.now().strftime("%Y%m%d_%H%M")}.csv')
        else:
            raise ValueError(f'Unknown format: {format}')


# === Helpers ===

def _format_text_table(headers, rows):
    """Format table as monospaced text."""
    # Convert everything to string
    str_rows = [[str(c) for c in row] for row in rows]
    str_headers = [str(h) for h in headers]

    # Column widths
    widths = [len(h) for h in str_headers]
    for row in str_rows:
        for i, c in enumerate(row):
            widths[i] = max(widths[i], len(c))

    def format_row(cells, widths, is_num_col=None):
        parts = []
        for i, c in enumerate(cells):
            # Right-align if numeric column
            if is_num_col and is_num_col[i]:
                parts.append(c.rjust(widths[i]))
            else:
                parts.append(c.ljust(widths[i]))
        return '  '.join(parts)

    # Detect numeric columns
    is_num = [True] * len(str_headers)
    for row in str_rows:
        for i, c in enumerate(row):
            try:
                float(c.replace(' ', '').replace('%', '').replace(',', '').replace('+', '').replace('-', ''))
            except ValueError:
                is_num[i] = False

    result = []
    result.append(format_row(str_headers, widths, is_num))
    result.append('─' * sum(widths + [2 * (len(widths) - 1)]))
    for row in str_rows:
        result.append(format_row(row, widths, is_num))
    return result


def _wrap_text(text, width):
    """Simple text wrap."""
    lines = []
    for para in text.split('\n'):
        if len(para) <= width:
            lines.append(para)
            continue
        words = para.split()
        current = []
        length = 0
        for w in words:
            if length + len(w) + 1 > width:
                lines.append(' '.join(current))
                current = [w]
                length = len(w)
            else:
                current.append(w)
                length += len(w) + 1
        if current:
            lines.append(' '.join(current))
    return lines


if __name__ == '__main__':
    # Demo
    r = Report(title='Демо-отчёт', period='2026-04-07 – 2026-05-13')
    r.section('Ключевые цифры')
    r.kv('Всего сделок', 261)
    r.kv('WON', 0, highlight='red', comment='Ни одной победы')
    r.kv('LOSE', 178, highlight='yellow')
    r.section('Таблица')
    r.table(['Метрика', 'Значение'], [['Дозвон', '80%'], ['Средний чек', '273K']])
    r.paragraph('Захаров работает конвейером — много звонков, ноль побед.', style='warning')
    r.conclusion(
        strong=['Много звонков', 'Хорошая дозваниваемость'],
        weak=['0 WON', 'Много LOSE'],
        recs=['Перевести на квалификацию', 'Дать ментора']
    )
    print(r.to_text())
